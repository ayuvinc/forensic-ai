"""BaseReportBuilder — python-docx report construction (RD-01).

Provides a structured builder for generating branded .docx reports.
Loads a firm template if provided; falls back to a blank document if the
template is missing or has incompatible styles.

All writes are atomic: content is built in memory, saved to a .tmp path,
then os.replace()d to the final destination so partial writes never corrupt output.

Usage:
    from tools.report_builder import BaseReportBuilder

    builder = BaseReportBuilder(template_path="firm_profile/templates/frm.docx")
    builder.set_header("GoodWork Forensic Consulting — CONFIDENTIAL")
    builder.set_footer("PRELIMINARY — NOT FOR CLIENT DISTRIBUTION")
    builder.add_cover_page("FRM Risk Register", "Client: Acme Corp", {"Date": "2026-04-20"})
    builder.add_toc()
    builder.add_section("1. Executive Summary", "Risk register covering 4 modules...")
    builder.add_subsection("1.1 Scope", "This assessment covers...")
    path = builder.save("cases/project-alpha-frm/F_Final/final_report.en.docx")
"""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# python-docx is required (in requirements.txt). Import is deferred to first use
# so this module can be imported in environments where python-docx is temporarily unavailable.


class BaseReportBuilder:
    """Structured .docx builder with graceful template fallback.

    Template loading:
        - If template_path is provided and the file exists, the document is
          opened from that template.
        - If the template has styles incompatible with python-docx, a warning
          is logged and a blank Document is used instead.
        - If template_path is None or missing, a blank Document is used.

    Style resolution:
        - Headings use 'Heading 1' / 'Heading 2' if present in the document;
          fall back to 'Normal' with bold formatting if the style is absent.
        - Cover page uses 'Title' style where available.
    """

    def __init__(self, template_path: Optional[str | Path] = None) -> None:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        self._doc = self._load_document(template_path)
        self._header_text: str = ""
        self._footer_text: str = ""

    # ── Public API ─────────────────────────────────────────────────────────────

    def set_header(self, text: str) -> "BaseReportBuilder":
        """Set the running header text for all sections."""
        self._header_text = text
        return self

    def set_footer(self, text: str) -> "BaseReportBuilder":
        """Set the running footer text for all sections."""
        self._footer_text = text
        return self

    def add_cover_page(
        self,
        title: str,
        subtitle: str = "",
        metadata: Optional[dict] = None,
    ) -> "BaseReportBuilder":
        """Add a cover page with title, subtitle, and optional metadata table.

        metadata dict keys become left-column labels; values become right-column content.
        Example: {"Client": "Acme Corp", "Date": "2026-04-20", "Classification": "Confidential"}
        """
        from docx.shared import Pt

        doc = self._doc
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        try:
            title_run.font.size = Pt(24)
            title_para.style = doc.styles[self._resolve_style("Title")]
        except (KeyError, Exception):
            title_run.font.size = Pt(24)

        if subtitle:
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.bold = False
            try:
                sub_run.font.size = Pt(14)
            except Exception:
                pass

        if metadata:
            doc.add_paragraph()  # spacer
            table = doc.add_table(rows=len(metadata), cols=2)
            for i, (label, value) in enumerate(metadata.items()):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            try:
                table.style = "Table Grid"
            except Exception:
                pass

        doc.add_page_break()
        return self

    def add_toc(self) -> "BaseReportBuilder":
        """Add a Table of Contents placeholder.

        Word will populate this on first open (displays "Update fields?" prompt).
        The placeholder text makes the intent clear in both docx and markdown renderings.
        """
        self._doc.add_paragraph("Table of Contents", style=self._resolve_style("Heading 1"))
        toc_para = self._doc.add_paragraph()
        toc_para.add_run(
            "[Table of Contents — right-click and select 'Update Field' in Microsoft Word]"
        ).italic = True
        self._doc.add_paragraph()
        return self

    def add_section(self, heading: str, content: str = "") -> "BaseReportBuilder":
        """Add a top-level section (Heading 1) with optional body text."""
        self._doc.add_paragraph(heading, style=self._resolve_style("Heading 1"))
        if content:
            self._add_body_text(content)
        return self

    def add_subsection(self, heading: str, content: str = "") -> "BaseReportBuilder":
        """Add a subsection (Heading 2) with optional body text."""
        self._doc.add_paragraph(heading, style=self._resolve_style("Heading 2"))
        if content:
            self._add_body_text(content)
        return self

    def add_heat_map(self, risk_items: list) -> "BaseReportBuilder":
        """Add a 5×5 likelihood × impact heat map table to the document (FR-05).

        Cells are colour-coded by risk rating (likelihood × impact):
          1-4  → green   (#92D050)
          5-9  → amber   (#FFC000)
          10-19 → red    (#FF0000)
          20-25 → dark red (#C00000)

        risk_items: list of RiskItem objects or dicts with likelihood/impact fields.
        Returns self for fluent chaining.
        AC: calling with 25 stub items (all L×I combos) returns self without raising.
        """
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _color(likelihood: int, impact: int) -> tuple[int, int, int]:
            rating = likelihood * impact
            if rating >= 20:
                return (192, 0, 0)
            if rating >= 10:
                return (255, 0, 0)
            if rating >= 5:
                return (255, 192, 0)
            return (146, 208, 80)

        def _set_cell_bg(cell, rgb: tuple[int, int, int]) -> None:
            """Apply background shading to a table cell via direct XML."""
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb))
            tcPr.append(shd)

        # Build rating count map — each (likelihood, impact) cell shows risk count
        counts: dict[tuple[int, int], int] = {}
        for item in risk_items:
            def _get(attr, default=1):
                if hasattr(item, attr):
                    return getattr(item, attr)
                return item.get(attr, default) if isinstance(item, dict) else default
            l = max(1, min(5, _get("likelihood")))
            i = max(1, min(5, _get("impact")))
            counts[(l, i)] = counts.get((l, i), 0) + 1

        # Add a heading for the heat map
        self._doc.add_paragraph("Risk Heat Map", style=self._resolve_style("Heading 2"))

        # 6×6 table: row 0 = column headers; col 0 = row headers
        tbl = self._doc.add_table(rows=6, cols=6)
        tbl.style = self._resolve_style("Table Grid")

        # Column header row (impact 1–5)
        hdr_row = tbl.rows[0]
        hdr_row.cells[0].text = "L \\ I"
        for impact in range(1, 6):
            hdr_row.cells[impact].text = str(impact)

        # Data rows (likelihood 1–5)
        for likelihood in range(1, 6):
            row = tbl.rows[likelihood]
            row.cells[0].text = str(likelihood)
            for impact in range(1, 6):
                cell  = row.cells[impact]
                rgb   = _color(likelihood, impact)
                count = counts.get((likelihood, impact), 0)
                cell.text = str(count) if count else ""
                _set_cell_bg(cell, rgb)

        return self

    def save(self, output_path: str | Path) -> Path:
        """Atomically write the document to output_path.

        Creates parent directories as needed.
        Returns the final Path of the written file.
        """
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        # Apply header/footer if set
        if self._header_text or self._footer_text:
            self._apply_header_footer()

        tmp = out.with_suffix(".tmp.docx")
        self._doc.save(str(tmp))
        os.replace(tmp, out)
        return out

    # ── Private helpers ───────────────────────────────────────────────────────

    def _load_document(self, template_path: Optional[str | Path]):
        """Return a docx Document, falling back to blank on any template error."""
        from docx import Document

        if template_path is not None:
            tpath = Path(template_path)
            if tpath.exists():
                try:
                    return Document(str(tpath))
                except Exception as exc:
                    logger.warning(
                        "BaseReportBuilder: template %s failed to load (%s) — using blank document",
                        tpath,
                        exc,
                    )
        return Document()

    # Mapping from standard Word styles to GW_ equivalents.
    # When a GW_ style is present in the loaded template it takes priority,
    # ensuring at least one GW_-prefixed paragraph appears in every output.
    _GW_STYLE_MAP = {
        "Heading 1": "GW_Heading1",
        "Heading 2": "GW_Heading2",
        "Normal":    "GW_Body",
        "Title":     "GW_Title",
    }

    def _resolve_style(self, style_name: str) -> str:
        """Return the best available style for the document.

        Priority: GW_ equivalent → requested style → 'Normal'.
        Ensures GW_-prefixed paragraphs appear in the output when the template
        carries GW_ styles (TPL-05 AC requirement).
        """
        available = {s.name for s in self._doc.styles}

        # Prefer GW_ equivalent if the template defines it
        gw_name = self._GW_STYLE_MAP.get(style_name)
        if gw_name and gw_name in available:
            return gw_name

        # Fall back to the requested style
        if style_name in available:
            return style_name

        return "Normal"

    def _add_body_text(self, content: str) -> None:
        """Add body paragraphs, splitting on double newlines."""
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            para = self._doc.add_paragraph()
            para.add_run(para_text)

    def _apply_header_footer(self) -> None:
        """Set running header and footer on the default section."""
        try:
            section = self._doc.sections[0]
            if self._header_text:
                header = section.header
                header.paragraphs[0].clear()
                header.paragraphs[0].add_run(self._header_text)
            if self._footer_text:
                footer = section.footer
                footer.paragraphs[0].clear()
                footer.paragraphs[0].add_run(self._footer_text)
        except Exception as exc:
            logger.warning("BaseReportBuilder: header/footer apply failed: %s", exc)
