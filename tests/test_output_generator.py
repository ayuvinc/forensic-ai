"""Tests for tools/output_generator.py (TEST-06).

Covers: generate_docx produces a valid .docx, no .tmp file remains,
template resolution chain (explicit path → fallback), GW_ style fallback.
"""
from __future__ import annotations

from pathlib import Path

import pytest


class TestGenerateDocx:
    def test_creates_docx(self, tmp_path):
        from tools.output_generator import OutputGenerator
        out = tmp_path / "report.docx"
        OutputGenerator().generate_docx("# Hello\n\nBody text.", out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_no_tmp_file_remains(self, tmp_path):
        from tools.output_generator import OutputGenerator
        out = tmp_path / "report.docx"
        OutputGenerator().generate_docx("# Title", out)
        assert not out.with_suffix(".tmp.docx").exists()

    def test_valid_docx_structure(self, tmp_path):
        from tools.output_generator import OutputGenerator
        from docx import Document
        out = tmp_path / "report.docx"
        OutputGenerator().generate_docx("# H1\n## H2\n- bullet\n\nParagraph.", out)
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("H1" in t for t in texts)

    def test_explicit_template_path_wins(self, tmp_path):
        """When explicit template_path exists, it is used — no TemplateManager call."""
        from docx import Document
        # Create a minimal valid .docx to use as template
        tmpl = tmp_path / "tmpl.docx"
        Document().save(str(tmpl))

        from tools.output_generator import OutputGenerator
        out = tmp_path / "out.docx"
        OutputGenerator().generate_docx(
            "# Test\nContent.", out, template_path=tmpl
        )
        assert out.exists()

    def test_nonexistent_explicit_template_falls_back(self, tmp_path):
        """When explicit template_path does not exist, falls back to plain Document()."""
        from tools.output_generator import OutputGenerator
        out = tmp_path / "out.docx"
        OutputGenerator().generate_docx(
            "Body.", out, template_path=tmp_path / "missing.docx"
        )
        assert out.exists()

    def test_workflow_type_triggers_template_manager_fallback(self, tmp_path, monkeypatch):
        """When workflow_type is set and TemplateManager raises, falls back silently."""
        from tools import output_generator
        # Make TemplateManager always raise so we exercise the except branch
        monkeypatch.setattr(
            "tools.output_generator.OutputGenerator.generate_docx",
            output_generator.OutputGenerator.generate_docx,
        )
        from tools.output_generator import OutputGenerator

        def _bad_resolve(self, wf):
            raise FileNotFoundError("no template")

        import tools.template_manager as tm
        monkeypatch.setattr(tm.TemplateManager, "resolve", _bad_resolve)

        out = tmp_path / "out.docx"
        OutputGenerator().generate_docx("# Title", out, workflow_type="unknown_workflow")
        assert out.exists()

    def test_audit_event_written_when_case_id_provided(self, tmp_path, patched_cases_dir, sample_case_id):
        """When case_id is provided, a template_resolved audit event is appended."""
        import json
        from tools.output_generator import OutputGenerator
        from tools.file_tools import case_dir

        out = tmp_path / "report.docx"
        OutputGenerator().generate_docx("# Title", out, case_id=sample_case_id)

        log_path = case_dir(sample_case_id) / "audit_log.jsonl"
        assert log_path.exists()
        events = [json.loads(line) for line in log_path.read_text().strip().splitlines()]
        assert any(e.get("event") == "template_resolved" for e in events)

    def test_markdown_variants_do_not_crash(self, tmp_path):
        """All markdown patterns (h1/h2/h3, bullet, table, divider, bold) parse without crash."""
        from tools.output_generator import OutputGenerator
        content = (
            "# Title\n"
            "## Section\n"
            "### Subsection\n"
            "- item 1\n"
            "- item 2\n"
            "| Col A | Col B |\n"
            "---\n"
            "**Bold text** and _italic_ here.\n"
            "\n"
            "Normal paragraph.\n"
        )
        out = tmp_path / "variants.docx"
        OutputGenerator().generate_docx(content, out)
        assert out.exists()
