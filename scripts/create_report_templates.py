"""
create_report_templates.py
Generates base Word (.docx) templates for GoodWork Forensic AI.
Run: python scripts/create_report_templates.py
"""

from pathlib import Path
import json
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ---------------------------------------------------------------------------
# Brand colours
# ---------------------------------------------------------------------------
GW_RED        = RGBColor(0xD5, 0x00, 0x32)
GW_DARK       = RGBColor(0x28, 0x28, 0x27)
GW_NEUTRAL    = RGBColor(0x4F, 0x4F, 0x4E)
GW_WARM       = RGBColor(0xF5, 0xF2, 0xF0)
GW_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GW_LIGHT_LINE = RGBColor(0xD5, 0xD5, 0xD5)

RAG_GREEN_TEXT = RGBColor(0x1B, 0x5E, 0x20)
RAG_GREEN_FILL = RGBColor(0xE8, 0xF5, 0xE9)
RAG_AMBER_TEXT = RGBColor(0xE6, 0x51, 0x00)
RAG_AMBER_FILL = RGBColor(0xFF, 0xF3, 0xE0)
RAG_RED_TEXT   = RGBColor(0xFF, 0xFF, 0xFF)
RAG_RED_FILL   = RGBColor(0xD5, 0x00, 0x32)

FONT_BODY    = "Calibri"
FONT_HEADING = "Calibri Light"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_cell_shading(cell, fill_color: RGBColor):
    """Set background fill on a table cell via <w:shd>."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(fill_color))
    # Remove existing shd if present
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_paragraph_border(para, side: str, color: RGBColor, sz: int = 24, val: str = "single"):
    """Add a border on one side of a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), _hex(color))
    # Remove existing border on this side
    existing = pBdr.find(qn(f"w:{side}"))
    if existing is not None:
        pBdr.remove(existing)
    pBdr.append(border)


def _set_full_paragraph_border(para, color: RGBColor, sides=("top", "bottom", "left", "right")):
    """Set borders on multiple sides of a paragraph."""
    for side in sides:
        _set_paragraph_border(para, side, color)


def _add_horizontal_rule_paragraph(doc, color: RGBColor = GW_RED) -> None:
    """Add an empty paragraph whose bottom border acts as a full-width rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")   # 1pt = 8 eighths-of-a-point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    # Zero spacing so the rule sits tight
    spacingEl = OxmlElement("w:spacing")
    spacingEl.set(qn("w:before"), "0")
    spacingEl.set(qn("w:after"), "0")
    pPr.append(spacingEl)


def _set_line_spacing(para, lines: float):
    """Set line spacing as a multiple (e.g. 1.15)."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    val = int(lines * 240)
    spacing.set(qn("w:line"), str(val))
    spacing.set(qn("w:lineRule"), "auto")


def _set_para_spacing(para, before_pt: int = 0, after_pt: int = 6):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before_pt * 20))
    spacing.set(qn("w:after"), str(after_pt * 20))


def _set_run_font(run, name: str, size_pt: float, bold: bool = False,
                  italic: bool = False, color: RGBColor = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Also set theme font override so Calibri renders correctly
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _set_cell_width(cell, width_inches: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


# ---------------------------------------------------------------------------
# Public helper functions
# ---------------------------------------------------------------------------

def set_paragraph_style(para, font_name: str, size_pt: float,
                         bold: bool = False, color: RGBColor = None,
                         space_before: int = 0, space_after: int = 6,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         italic: bool = False) -> None:
    """Apply font/spacing to a paragraph's default run format."""
    para.alignment = alignment
    _set_para_spacing(para, space_before, space_after)
    # Apply to all existing runs
    for run in para.runs:
        _set_run_font(run, font_name, size_pt, bold, italic, color)
    # Also set paragraph-level rPr defaults via pPr/rPr so new runs inherit
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    # font
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    # size
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    # bold
    b_el = rPr.find(qn("w:b"))
    if bold:
        if b_el is None:
            rPr.append(OxmlElement("w:b"))
    else:
        if b_el is not None:
            rPr.remove(b_el)
    # color
    if color:
        c_el = rPr.find(qn("w:color"))
        if c_el is None:
            c_el = OxmlElement("w:color")
            rPr.append(c_el)
        c_el.set(qn("w:val"), _hex(color))


def add_red_rule(doc: Document) -> None:
    """Add a full-width horizontal rule in GW_RED."""
    _add_horizontal_rule_paragraph(doc, GW_RED)


def add_cover_page(doc: Document, title: str, subtitle: str,
                   client_name: str, date_str: str,
                   confidential: bool = True,
                   confidential_label: str = "CONFIDENTIAL") -> None:
    """Big-4 style cover page."""
    # Top rule
    add_red_rule(doc)

    # Firm name
    p_firm = doc.add_paragraph()
    _set_para_spacing(p_firm, 6, 0)
    r = p_firm.add_run("GoodWork Forensic Consulting")
    _set_run_font(r, FONT_HEADING, 14, bold=False, color=GW_NEUTRAL)

    # Large white space — 8 empty lines
    for _ in range(8):
        ep = doc.add_paragraph()
        _set_para_spacing(ep, 0, 0)

    # Document title
    p_title = doc.add_paragraph()
    _set_para_spacing(p_title, 0, 6)
    r_title = p_title.add_run(title)
    _set_run_font(r_title, FONT_HEADING, 28, bold=True, color=GW_DARK)

    # Subtitle
    p_sub = doc.add_paragraph()
    _set_para_spacing(p_sub, 0, 12)
    r_sub = p_sub.add_run(subtitle)
    _set_run_font(r_sub, FONT_HEADING, 16, bold=False, color=GW_RED)

    # Client
    p_client = doc.add_paragraph()
    _set_para_spacing(p_client, 0, 4)
    r_client = p_client.add_run(f"Prepared for: {client_name}")
    _set_run_font(r_client, FONT_BODY, 11, color=GW_NEUTRAL)

    # Date
    p_date = doc.add_paragraph()
    _set_para_spacing(p_date, 0, 4)
    r_date = p_date.add_run(date_str)
    _set_run_font(r_date, FONT_BODY, 11, color=GW_NEUTRAL)

    # Bottom rule
    add_red_rule(doc)

    # Confidential label
    if confidential:
        p_conf = doc.add_paragraph()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_spacing(p_conf, 6, 0)
        r_conf = p_conf.add_run(confidential_label)
        _set_run_font(r_conf, FONT_BODY, 9, bold=False, color=GW_RED)

    # Page break
    doc.add_page_break()


def add_header_footer(doc: Document,
                      firm_name: str = "GoodWork Forensic Consulting") -> None:
    """
    Header: firm name (left) | document title placeholder (right) — 9pt Calibri, GW_NEUTRAL,
            bottom border line.
    Footer: "Confidential | GoodWork Forensic Consulting | Page X" — 9pt, centered.
    """
    section = doc.sections[0]

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    # Clear default content
    for para in header.paragraphs:
        for run in para.runs:
            run.text = ""
    if not header.paragraphs:
        header.add_paragraph()

    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(hp, 0, 2)

    # Firm name (left)
    r_left = hp.add_run(firm_name)
    _set_run_font(r_left, FONT_BODY, 9, color=GW_NEUTRAL)

    # Tab to right-align document title
    r_tab = hp.add_run("\t")
    _set_run_font(r_tab, FONT_BODY, 9, color=GW_NEUTRAL)

    r_right = hp.add_run("[Document Title]")
    _set_run_font(r_right, FONT_BODY, 9, color=GW_NEUTRAL)

    # Set two tab stops: center + right
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "9350")  # ~6.5 inches from left margin
    tabs.append(tab_right)
    pPr.append(tabs)

    # Bottom border on header paragraph
    _set_paragraph_border(hp, "bottom", GW_LIGHT_LINE, sz=4)

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""
    if not footer.paragraphs:
        footer.add_paragraph()

    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(fp, 2, 0)

    r_conf = fp.add_run("Confidential  |  GoodWork Forensic Consulting  |  Page ")
    _set_run_font(r_conf, FONT_BODY, 9, color=GW_NEUTRAL)

    # Page number field
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r_page = fp.add_run()
    r_page._r.append(fldChar_begin)
    r_page._r.append(instrText)
    r_page._r.append(fldChar_end)
    _set_run_font(r_page, FONT_BODY, 9, color=GW_NEUTRAL)


def _get_or_add_style(doc: Document, style_name: str, style_type_str: str = "paragraph"):
    """Return existing named style or create a new one."""
    from docx.enum.style import WD_STYLE_TYPE
    style_type = WD_STYLE_TYPE.PARAGRAPH if style_type_str == "paragraph" else WD_STYLE_TYPE.CHARACTER
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, style_type)


def set_named_styles(doc: Document) -> None:
    """Define all GW_ named paragraph styles."""

    def _apply(style_name, font_name, size_pt, bold=False, italic=False,
               color=None, space_before=0, space_after=6,
               line_spacing=None, bottom_border=False, left_border=False,
               indent_left=None):
        style = _get_or_add_style(doc, style_name)
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.italic = italic
        if color:
            style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        if indent_left is not None:
            pf.left_indent = Inches(indent_left)
        if line_spacing is not None:
            from docx.shared import Length
            pf.line_spacing = Pt(size_pt * line_spacing)

        # Borders via XML
        pPr = style.element.get_or_add_pPr()
        if bottom_border or left_border:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = OxmlElement("w:pBdr")
                pPr.append(pBdr)
            if bottom_border:
                b = OxmlElement("w:bottom")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "8")
                b.set(qn("w:space"), "1")
                b.set(qn("w:color"), _hex(GW_RED))
                old = pBdr.find(qn("w:bottom"))
                if old is not None:
                    pBdr.remove(old)
                pBdr.append(b)
            if left_border:
                lb = OxmlElement("w:left")
                lb.set(qn("w:val"), "single")
                lb.set(qn("w:sz"), "24")   # 3pt
                lb.set(qn("w:space"), "4")
                lb.set(qn("w:color"), _hex(GW_RED))
                old = pBdr.find(qn("w:left"))
                if old is not None:
                    pBdr.remove(old)
                pBdr.append(lb)

    _apply("GW_Title",      FONT_HEADING, 28, bold=True,  color=GW_DARK,    space_before=0,  space_after=12)
    _apply("GW_Heading1",   FONT_HEADING, 16, bold=True,  color=GW_DARK,    space_before=18, space_after=6,  bottom_border=True)
    _apply("GW_Heading2",   FONT_HEADING, 13, bold=False, color=GW_RED,     space_before=12, space_after=4)
    _apply("GW_Heading3",   FONT_BODY,    11, bold=True,  color=GW_DARK,    space_before=8,  space_after=4)
    _apply("GW_Body",       FONT_BODY,    11, bold=False, color=GW_NEUTRAL, space_before=0,  space_after=6,  line_spacing=1.15)
    _apply("GW_TableHeader",FONT_BODY,    10, bold=True,  color=GW_WHITE,   space_before=0,  space_after=0)
    _apply("GW_Caption",    FONT_BODY,    9,  bold=False, color=GW_NEUTRAL, space_before=0,  space_after=8,  italic=True)
    _apply("GW_Disclaimer", FONT_BODY,    9,  bold=False, color=GW_NEUTRAL, space_before=0,  space_after=6,  italic=True, left_border=True, indent_left=0.3)
    _apply("GW_BASELINE",   FONT_BODY,    9,  bold=False, color=GW_NEUTRAL, space_before=0,  space_after=0,  italic=True)


# ---------------------------------------------------------------------------
# Utility: add styled paragraph
# ---------------------------------------------------------------------------

def _add_styled_para(doc: Document, text: str, style_name: str) -> None:
    """Add paragraph using a named style defined by set_named_styles."""
    p = doc.add_paragraph(style=style_name)
    p.add_run(text)


def _add_sample_sections(doc: Document) -> None:
    """Insert 2–3 sample sections demonstrating each named style."""
    _add_styled_para(doc, "1.  Executive Summary", "GW_Heading1")
    _add_styled_para(doc,
        "This document was generated by GoodWork Forensic AI. Replace this placeholder "
        "content with the actual executive summary for your engagement.",
        "GW_Body")

    _add_styled_para(doc, "1.1  Background", "GW_Heading2")
    _add_styled_para(doc,
        "Provide background context here. Describe the client, the circumstances leading "
        "to this engagement, and the key stakeholders involved.",
        "GW_Body")

    _add_styled_para(doc, "Key Finding", "GW_Heading3")
    _add_styled_para(doc,
        "Detail the finding with supporting evidence. Each finding should reference the "
        "specific documents reviewed and the methodology applied.",
        "GW_Body")

    _add_styled_para(doc,
        "Figure 1: Overview of engagement scope and key deliverables.",
        "GW_Caption")

    _add_styled_para(doc,
        "DISCLAIMER: This report has been prepared solely for the use of the client named "
        "above. GoodWork Forensic Consulting accepts no liability to any third party who "
        "may rely on this report.",
        "GW_Disclaimer")

    _add_styled_para(doc, "2.  Scope & Methodology", "GW_Heading1")
    _add_styled_para(doc,
        "Describe the scope of the engagement, including inclusions, exclusions, and any "
        "limitations on the work performed.",
        "GW_Body")

    _add_styled_para(doc, "2.1  Methodology", "GW_Heading2")
    _add_styled_para(doc,
        "Describe the analytical approach: document review, interviews conducted, data "
        "analytics procedures, regulatory lookups, and any forensic software used.",
        "GW_Body")

    _add_styled_para(doc, "3.  Findings", "GW_Heading1")
    _add_styled_para(doc,
        "Each finding below follows the structure: Procedure → Finding → Implication → "
        "Conclusion, in accordance with ACFE reporting standards.",
        "GW_Body")

    # BASELINE label example
    p_bl = doc.add_paragraph(style="GW_BASELINE")
    p_bl.add_run("BASELINE — this section was AI-generated and requires partner review.")


# ---------------------------------------------------------------------------
# create_base_template
# ---------------------------------------------------------------------------

def create_base_template(output_path: Path, doc_type: str,
                         sample_content: bool = True,
                         cover_title: str = None,
                         cover_subtitle: str = None,
                         confidential_label: str = "CONFIDENTIAL") -> None:
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)

    if sample_content:
        title    = cover_title    or f"[{doc_type.replace('_', ' ').title()}]"
        subtitle = cover_subtitle or "GoodWork Forensic Consulting"
        add_cover_page(doc, title=title, subtitle=subtitle,
                       client_name="[Client Name]",
                       date_str="[Date]",
                       confidential=True,
                       confidential_label=confidential_label)
        _add_sample_sections(doc)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Risk Register template
# ---------------------------------------------------------------------------

def _add_risk_table(doc: Document) -> None:
    _add_styled_para(doc, "Risk Register", "GW_Heading1")
    _add_styled_para(doc,
        "The table below presents identified fraud risks with RAG-rated severity. "
        "Ratings: Green 1–8 (low), Amber 9–15 (medium), Red 16–25 (high/critical).",
        "GW_Body")

    col_widths = [0.4, 2.2, 1.2, 0.9, 0.9, 0.9, 1.1, 0.9]
    headers    = ["#", "Risk Title", "Module", "Likelihood\n(1–5)", "Impact\n(1–5)",
                  "Rating\n(1–25)", "Owner", "Status"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, headers, col_widths)):
        _set_cell_shading(cell, GW_DARK)
        _set_cell_width(cell, w)
        p = cell.paragraphs[0]
        p.clear()
        p.style = doc.styles["GW_TableHeader"] if "GW_TableHeader" in [s.name for s in doc.styles] else p.style
        r = p.add_run(hdr)
        _set_run_font(r, FONT_BODY, 10, bold=True, color=GW_WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sample data rows: (risk_title, module, likelihood, impact, owner, status, rag)
    sample_rows = [
        ("Vendor Collusion Risk", "Module 3 — Procurement", 2, 4, "CFO",       "Open",  "green"),
        ("Payroll Ghost Employee", "Module 4 — HR & Payroll", 3, 3, "HR Dir.",   "Open",  "amber"),
        ("Management Override",   "Module 1 — Governance",   3, 5, "Board",     "Open",  "red"),
        ("Expense Fraud",         "Module 5 — Expenses",     4, 3, "Finance",   "Review","amber"),
        ("Bribery — Contracts",   "Module 6 — Third Parties",3, 5, "Legal",     "Open",  "red"),
        ("Data Theft",            "Module 7 — IT & Data",    2, 3, "IT Mgr.",   "Open",  "green"),
    ]

    rag_map = {
        "green": (RAG_GREEN_FILL, RAG_GREEN_TEXT),
        "amber": (RAG_AMBER_FILL, RAG_AMBER_TEXT),
        "red":   (RAG_RED_FILL,   RAG_RED_TEXT),
    }

    for idx, (rtitle, module, lik, imp, owner, status, rag) in enumerate(sample_rows, 1):
        rating = lik * imp
        fill_color, text_color = rag_map[rag]
        row = table.add_row()
        values = [str(idx), rtitle, module, str(lik), str(imp), str(rating), owner, status]

        for col_i, (cell, val, w) in enumerate(zip(row.cells, values, col_widths)):
            _set_cell_width(cell, w)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            # Rating column gets RAG fill + coloured text; others get plain white
            if col_i == 5:  # rating column
                _set_cell_shading(cell, fill_color)
                _set_run_font(r, FONT_BODY, 10, color=text_color)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _set_run_font(r, FONT_BODY, 10, color=GW_DARK)
                if col_i in (0, 3, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer


def create_frm_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Fraud Risk Register",
                   subtitle="Enterprise Fraud Risk Assessment",
                   client_name="[Client Name]",
                   date_str="[Date]")
    _add_sample_sections(doc)
    _add_risk_table(doc)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Investigation Report template
# ---------------------------------------------------------------------------

def _add_finding_card(doc: Document, finding_no: int, title: str,
                      is_critical: bool = False) -> None:
    """Add a finding block. Critical findings get a red left border."""
    style_name = "GW_Disclaimer" if is_critical else "GW_Heading3"
    label = f"Finding {finding_no}: {title}"
    p = doc.add_paragraph(style=style_name)
    r = p.add_run(label)
    if is_critical:
        _set_run_font(r, FONT_BODY, 11, bold=True, color=GW_DARK)
    _add_styled_para(doc,
        "Procedure: Describe the procedure applied to identify this finding.\n"
        "Finding: State the factual finding without opinion.\n"
        "Implication: Explain what this finding means in the context of the investigation.\n"
        "Conclusion: Set out the conclusion drawn and its basis.",
        "GW_Body")


def create_investigation_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Investigation Report",
                   subtitle="Confidential — Prepared for Management",
                   client_name="[Client Name]",
                   date_str="[Date]")

    sections_content = [
        ("1.  Executive Summary",
         "Provide a concise summary of the investigation, key findings, and recommended actions. "
         "This section is intended for senior management and the Board."),
        ("2.  Scope & Methodology",
         "Describe the scope of the investigation including dates, entities reviewed, "
         "and any agreed limitations. Set out the methodology applied."),
        ("3.  Background",
         "Provide context: how the matter came to light, initial allegations or triggers, "
         "and relevant organisational background."),
    ]
    for sec_title, sec_body in sections_content:
        _add_styled_para(doc, sec_title, "GW_Heading1")
        _add_styled_para(doc, sec_body, "GW_Body")

    _add_styled_para(doc, "4.  Findings", "GW_Heading1")
    _add_styled_para(doc,
        "Findings are presented in order of materiality. Critical findings are marked "
        "with a red indicator. All findings follow ACFE reporting standards.",
        "GW_Body")

    _add_finding_card(doc, 1, "Misappropriation of Company Assets", is_critical=True)
    _add_finding_card(doc, 2, "Unsupported Journal Entries — Q3 FY2024", is_critical=True)
    _add_finding_card(doc, 3, "Gaps in Internal Control Documentation", is_critical=False)

    _add_styled_para(doc, "5.  Evidence Chain", "GW_Heading1")
    _add_styled_para(doc,
        "The following evidence was obtained, reviewed, and relied upon in reaching "
        "the findings set out above. Evidence is listed in chronological order.",
        "GW_Body")
    _add_styled_para(doc, "Evidence Chain", "GW_Heading2")
    _add_styled_para(doc,
        "1. [Document name] — [date obtained] — [relevance to finding]\n"
        "2. [Interview transcript] — [interviewee] — [date] — [key statements]\n"
        "3. [Transaction data] — [period] — [source system]",
        "GW_Body")

    _add_styled_para(doc, "6.  Conclusions", "GW_Heading1")
    _add_styled_para(doc,
        "Set out the overall conclusions of the investigation based on the findings "
        "and evidence reviewed. Do not include opinions not supported by evidence.",
        "GW_Body")

    _add_styled_para(doc, "7.  Recommendations", "GW_Heading1")
    _add_styled_para(doc,
        "Set out remediation actions recommended to management, prioritised by urgency. "
        "Each recommendation should reference the finding it addresses.",
        "GW_Body")

    _add_styled_para(doc,
        "DISCLAIMER: This report has been prepared solely for the use of the named client. "
        "GoodWork Forensic Consulting accepts no liability to third parties.",
        "GW_Disclaimer")

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Client Proposal template
# ---------------------------------------------------------------------------

def _add_fee_table(doc: Document) -> None:
    _add_styled_para(doc, "5.  Fees", "GW_Heading1")
    _add_styled_para(doc,
        "Our proposed fees are set out below. All amounts are exclusive of VAT.",
        "GW_Body")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    col_widths = [4.5, 2.0]
    headers = ["Service / Deliverable", "Fee (USD)"]

    hdr_row = table.rows[0]
    for cell, hdr, w in zip(hdr_row.cells, headers, col_widths):
        _set_cell_shading(cell, GW_RED)
        _set_cell_width(cell, w)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(hdr)
        _set_run_font(r, FONT_BODY, 10, bold=True, color=GW_WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sample_fees = [
        ("Phase 1 — Document Review & Initial Analysis", "USD 15,000"),
        ("Phase 2 — Fieldwork & Interviews",              "USD 25,000"),
        ("Phase 3 — Reporting & Partner Review",          "USD 10,000"),
        ("Total",                                          "USD 50,000"),
    ]
    for svc, fee in sample_fees:
        row = table.add_row()
        for cell, val, w in zip(row.cells, [svc, fee], col_widths):
            _set_cell_width(cell, w)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            is_total = (svc == "Total")
            _set_run_font(r, FONT_BODY, 10, bold=is_total, color=GW_DARK)
            if is_total:
                _set_cell_shading(cell, GW_WARM)

    doc.add_paragraph()


def create_proposal_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Proposal for Forensic Consulting Services",
                   subtitle="[Engagement Description]",
                   client_name="[Prospective Client Name]",
                   date_str="[Date]",
                   confidential=True,
                   confidential_label="STRICTLY PRIVATE & CONFIDENTIAL")

    proposal_sections = [
        ("1.  Executive Summary",
         "Provide a concise summary of our understanding of the client's needs and how "
         "GoodWork Forensic Consulting is uniquely positioned to address them."),
        ("2.  Our Understanding",
         "Describe our understanding of the client's situation, the business context, "
         "and the key issues driving this engagement."),
        ("3.  Scope of Work",
         "Set out the precise scope of work: what is included, what is excluded, and "
         "the key deliverables to be provided."),
        ("4.  Methodology",
         "Describe the approach to be taken: phases of work, analytical methods, "
         "regulatory frameworks applied, and quality control procedures."),
    ]
    for sec_title, sec_body in proposal_sections:
        _add_styled_para(doc, sec_title, "GW_Heading1")
        _add_styled_para(doc, sec_body, "GW_Body")

    _add_fee_table(doc)

    _add_styled_para(doc, "6.  Our Team", "GW_Heading1")
    _add_styled_para(doc,
        "The following team members will be responsible for delivering this engagement.",
        "GW_Body")
    _add_styled_para(doc, "[Partner Name] — Engagement Partner", "GW_Heading3")
    _add_styled_para(doc, "Brief bio and relevant experience.", "GW_Body")

    _add_styled_para(doc, "7.  Terms & Conditions", "GW_Heading1")
    _add_styled_para(doc,
        "This proposal is subject to GoodWork Forensic Consulting's standard terms and "
        "conditions. A copy is appended to this proposal.",
        "GW_Body")

    _add_styled_para(doc,
        "STRICTLY PRIVATE & CONFIDENTIAL: This proposal has been prepared exclusively for "
        "the named recipient and may not be shared or disclosed without written consent.",
        "GW_Disclaimer")

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Workpaper template
# ---------------------------------------------------------------------------

def _add_workpaper_banner(doc: Document) -> None:
    """Red-background full-width banner: INTERIM WORKPAPER — NOT FOR CLIENT DISTRIBUTION."""
    # Use a 1-row 1-cell table to get a solid red background spanning the full width
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    # Full page width minus margins
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, GW_RED)
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, 6, 6)
    r = p.add_run("INTERIM WORKPAPER — NOT FOR CLIENT DISTRIBUTION")
    _set_run_font(r, FONT_BODY, 11, bold=True, color=GW_WHITE)
    doc.add_paragraph()  # spacer after banner


def _add_draft_watermark(doc: Document) -> None:
    """
    Insert a DRAFT diagonal text watermark via DrawingML in the header.
    This is the standard Word watermark approach.
    """
    section = doc.sections[0]
    header = section.header

    # We add the watermark XML to the first header paragraph
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]

    # Build the watermark XML fragment
    WATERMARK_XML = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '<w:pPr><w:rPr><w:rStyle w:val="PlaceholderText"/></w:rPr></w:pPr>'
        '<w:r><w:rPr><w:rStyle w:val="PlaceholderText"/></w:rPr>'
        '<w:pict><v:shape id="_x0000_i1025" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:528pt;height:132pt;'
        'z-index:-251657216;mso-position-horizontal:center;mso-position-vertical:center" '
        'fillcolor="#D50032" stroked="f">'
        '<v:fill opacity=".3"/>'
        '<v:textpath style="font-family:\'Calibri\';font-size:72pt;font-weight:bold;" '
        'string="DRAFT" trim="t" fitshape="t"/>'
        '</v:shape></w:pict></w:r></w:p>'
    )

    from lxml import etree
    try:
        wm_el = etree.fromstring(WATERMARK_XML)
        hp._p.addnext(wm_el)
    except Exception:
        # Fallback: add text note in header if XML fails
        r_wm = hp.add_run(" [DRAFT]")
        _set_run_font(r_wm, FONT_BODY, 9, bold=True, color=GW_RED)


def create_workpaper_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc, firm_name="GoodWork Forensic Consulting — DRAFT")

    # Draft watermark
    _add_draft_watermark(doc)

    # Banner
    _add_workpaper_banner(doc)

    # Cover-style header (no page break)
    add_red_rule(doc)
    p_firm = doc.add_paragraph()
    r = p_firm.add_run("GoodWork Forensic Consulting")
    _set_run_font(r, FONT_HEADING, 14, color=GW_NEUTRAL)
    _set_para_spacing(p_firm, 6, 4)

    p_title = doc.add_paragraph()
    r2 = p_title.add_run("Interim Workpaper")
    _set_run_font(r2, FONT_HEADING, 22, bold=True, color=GW_DARK)
    _set_para_spacing(p_title, 0, 4)

    p_sub = doc.add_paragraph()
    r3 = p_sub.add_run("[Workpaper Title / Procedure Reference]")
    _set_run_font(r3, FONT_HEADING, 14, color=GW_RED)
    _set_para_spacing(p_sub, 0, 8)

    p_meta = doc.add_paragraph()
    r4 = p_meta.add_run("Case: [Case ID]  |  Prepared by: [Name]  |  Date: [Date]  |  Status: DRAFT")
    _set_run_font(r4, FONT_BODY, 10, color=GW_NEUTRAL)
    _set_para_spacing(p_meta, 0, 4)

    add_red_rule(doc)
    doc.add_paragraph()

    _add_styled_para(doc, "Objective", "GW_Heading1")
    _add_styled_para(doc, "State the objective of this workpaper and the specific procedure being documented.", "GW_Body")

    _add_styled_para(doc, "Procedure Performed", "GW_Heading1")
    _add_styled_para(doc, "Describe the procedure in detail including the source data, period covered, and methodology applied.", "GW_Body")

    _add_styled_para(doc, "Findings & Observations", "GW_Heading1")
    _add_styled_para(doc, "Document all findings and observations. Flag items requiring follow-up with [FOLLOW-UP].", "GW_Body")

    _add_styled_para(doc, "Conclusion", "GW_Heading1")
    _add_styled_para(doc, "State the conclusion reached from this procedure. Reference any findings that will be carried into the main report.", "GW_Body")

    _add_styled_para(doc,
        "WORKPAPER NOTICE: This document is an interim workpaper prepared for internal "
        "professional use only. It is not a client deliverable and must not be shared "
        "with the client or any third party without partner approval.",
        "GW_Disclaimer")

    p_bl = doc.add_paragraph(style="GW_BASELINE")
    p_bl.add_run("BASELINE — AI-generated workpaper template. Requires preparer review before use.")

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Specialised templates that reuse create_base_template
# ---------------------------------------------------------------------------

def create_dd_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Due Diligence Report",
                   subtitle="Confidential — Subject Profile & Findings",
                   client_name="[Client Name]",
                   date_str="[Date]")

    dd_sections = [
        ("1.  Executive Summary", "Summary of due diligence findings and overall risk assessment."),
        ("2.  Subject Profile",
         "Full legal name, registered address, date of incorporation, directors, "
         "UBOs, and corporate structure."),
        ("3.  Regulatory & Licensing Status",
         "Regulatory licences held, jurisdictions of operation, and compliance history."),
        ("4.  Adverse Media & Reputational Risk",
         "Summary of adverse media findings across all jurisdictions reviewed."),
        ("5.  Sanctions & PEP Screening",
         "Results of sanctions screening (OFAC, UN, EU) and PEP check."),
        ("6.  Financial Overview",
         "Key financial indicators sourced from public filings and credit references."),
        ("7.  Findings & Risk Rating",
         "Overall risk rating (Low / Medium / High) with supporting rationale."),
        ("8.  Recommendations",
         "Recommended next steps, enhanced due diligence requirements, or conditions."),
    ]
    for sec_title, sec_body in dd_sections:
        _add_styled_para(doc, sec_title, "GW_Heading1")
        _add_styled_para(doc, sec_body, "GW_Body")

    _add_styled_para(doc,
        "DISCLAIMER: This report is based on publicly available information and "
        "authoritative databases as at the date of preparation. GoodWork Forensic "
        "Consulting does not warrant the completeness of third-party data sources.",
        "GW_Disclaimer")

    doc.save(str(output_path))


def create_sanctions_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Sanctions Screening Memorandum",
                   subtitle="List Results — Authoritative Sources Only",
                   client_name="[Client Name]",
                   date_str="[Date]")

    sanc_sections = [
        ("1.  Purpose & Scope",
         "This memorandum documents the results of sanctions screening conducted "
         "against authoritative lists: OFAC SDN, UN Security Council Consolidated List, "
         "and EU Financial Sanctions Database."),
        ("2.  Subjects Screened",
         "List all individuals and entities screened, including name variations and aliases."),
        ("3.  Screening Results",
         "For each subject: list searched, search date, match status, and disposition."),
        ("4.  Potential Matches — Analysis",
         "For any potential matches: name similarity analysis, date of birth / "
         "incorporation check, nationality / jurisdiction check, and final disposition."),
        ("5.  Conclusion",
         "Overall conclusion: clear / potential match requiring escalation / "
         "confirmed match. Escalation recommended where applicable."),
        ("6.  Limitations",
         "Screening is conducted against lists as at the date of this memorandum. "
         "Ongoing monitoring is recommended for higher-risk relationships."),
    ]
    for sec_title, sec_body in sanc_sections:
        _add_styled_para(doc, sec_title, "GW_Heading1")
        _add_styled_para(doc, sec_body, "GW_Body")

    _add_styled_para(doc,
        "AUTHORITATIVE SOURCES ONLY: GoodWork Forensic Consulting relies exclusively on "
        "OFAC (ofac.treas.gov), UN Security Council (un.org/securitycouncil), and EU "
        "(sanctions.ec.europa.eu) for sanctions determinations. No inference from "
        "secondary sources.",
        "GW_Disclaimer")

    doc.save(str(output_path))


def create_transaction_testing_template(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    set_named_styles(doc)
    add_header_footer(doc)
    add_cover_page(doc,
                   title="Transaction Testing Report",
                   subtitle="Test Plan, Results & Findings",
                   client_name="[Client Name]",
                   date_str="[Date]")

    tt_sections = [
        ("1.  Executive Summary",
         "Summary of transaction testing scope, key findings, and exceptions identified."),
        ("2.  Testing Scope & Objectives",
         "Define the population tested, sampling methodology, period covered, and "
         "testing objectives for each procedure."),
        ("3.  Test Plan",
         "Detail each test procedure: population, sample size, selection method, "
         "threshold, and pass/fail criteria."),
        ("4.  Results by Test Procedure",
         "For each test: number tested, number of exceptions, exception rate, "
         "and disposition of exceptions."),
        ("5.  Exceptions Log",
         "Full detail of each exception: transaction reference, date, amount, "
         "nature of exception, and risk classification."),
        ("6.  Extrapolation & Population Risk",
         "Where applicable, extrapolation of exception rate to the full population "
         "and estimated financial exposure."),
        ("7.  Conclusions & Recommendations",
         "Overall conclusions and recommended control improvements."),
    ]
    for sec_title, sec_body in tt_sections:
        _add_styled_para(doc, sec_title, "GW_Heading1")
        _add_styled_para(doc, sec_body, "GW_Body")

    _add_styled_para(doc,
        "DISCLAIMER: Transaction testing results are limited to the sample selected. "
        "Findings may not be representative of the full population unless extrapolation "
        "is explicitly performed and documented herein.",
        "GW_Disclaimer")

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Master entry point
# ---------------------------------------------------------------------------

# Base templates go to assets/templates/ (committed). Custom overrides go to firm_profile/templates/ (gitignored).
TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "assets" / "templates"

TEMPLATES_METADATA = {
    "base_report": {
        "base": "base_report_base.docx",
        "custom": None,
        "description": "General report — all workflows"
    },
    "frm_risk_register": {
        "base": "frm_risk_register_base.docx",
        "custom": None,
        "description": "FRM Risk Register — RAG risk table"
    },
    "investigation_report": {
        "base": "investigation_report_base.docx",
        "custom": None,
        "description": "Investigation Report — evidence chain structure"
    },
    "client_proposal": {
        "base": "client_proposal_base.docx",
        "custom": None,
        "description": "Client Proposal — 7-section commercial format"
    },
    "due_diligence": {
        "base": "due_diligence_base.docx",
        "custom": None,
        "description": "Due Diligence Report — subject profile + findings"
    },
    "sanctions_screening": {
        "base": "sanctions_screening_base.docx",
        "custom": None,
        "description": "Sanctions Screening Memo — list results format"
    },
    "transaction_testing": {
        "base": "transaction_testing_base.docx",
        "custom": None,
        "description": "Transaction Testing Report — test plan + results"
    },
    "workpaper": {
        "base": "workpaper_base.docx",
        "custom": None,
        "description": "Interim Workpaper — NOT FOR CLIENT DISTRIBUTION"
    },
}


def create_all_templates() -> None:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    generators = [
        ("base_report_base.docx",
         lambda p: create_base_template(p, "base_report", sample_content=True)),
        ("frm_risk_register_base.docx",
         create_frm_template),
        ("investigation_report_base.docx",
         create_investigation_template),
        ("client_proposal_base.docx",
         create_proposal_template),
        ("due_diligence_base.docx",
         create_dd_template),
        ("sanctions_screening_base.docx",
         create_sanctions_template),
        ("transaction_testing_base.docx",
         create_transaction_testing_template),
        ("workpaper_base.docx",
         create_workpaper_template),
    ]

    for filename, generator in generators:
        path = TEMPLATES_DIR / filename
        try:
            generator(path)
            print(f"Created: {path}")
        except Exception as exc:
            print(f"ERROR creating {filename}: {exc}")
            raise

    # Write metadata JSON
    metadata_path = TEMPLATES_DIR / "templates.json"
    metadata_path.write_text(
        json.dumps(TEMPLATES_METADATA, indent=2, ensure_ascii=False),
        encoding="utf-8"
    )
    print(f"Created: {metadata_path}")
    print(f"\nAll templates written to: {TEMPLATES_DIR}")


if __name__ == "__main__":
    create_all_templates()
