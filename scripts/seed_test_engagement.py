"""Seed script — creates a test engagement with 3 completed workflow cases.

Run once before testing Sprint-IA-01 IA-04 (Workspace multi-workflow outputs).
Creates:
  cases/abc-corp-test-engagement/           (AF project)
  cases/case_test_001/                      (investigation_report — DELIVERABLE_WRITTEN)
  cases/case_test_002/                      (frm_risk_register — DELIVERABLE_WRITTEN)
  cases/case_test_003/                      (due_diligence — DELIVERABLE_WRITTEN)

No API calls. Safe to run repeatedly (skips if engagement already exists).
"""

from __future__ import annotations

import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path

# Ensure project root is on the path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import CASES_DIR
from schemas.project import ProjectIntake
from tools.project_manager import AF_FOLDERS, ProjectManager

SLUG            = "abc-corp-test-engagement"
CLIENT_NAME     = "ABC Corp"
PROJECT_NAME    = "ABC Corp Test Engagement"
CASE_IDS        = {
    "investigation_report": "case_test_001",
    "frm_risk_register":    "case_test_002",
    "due_diligence":        "case_test_003",
}
NOW = datetime.now(timezone.utc).isoformat()


def _write_json(path: Path, data: object) -> None:
    tmp = path.with_suffix(".tmp")
    tmp.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    os.replace(tmp, path)


def _seed_case(workflow_type: str, case_id: str) -> None:
    cdir = CASES_DIR / case_id
    cdir.mkdir(parents=True, exist_ok=True)
    (cdir / "F_Final").mkdir(exist_ok=True)

    _write_json(cdir / "state.json", {
        "case_id":      case_id,
        "workflow":     workflow_type,
        "status":       "deliverable_written",
        "last_updated": NOW,
    })

    placeholder = (
        f"# {workflow_type.replace('_', ' ').title()} — Test Report\n\n"
        f"Client: {CLIENT_NAME}\n"
        f"Case ID: {case_id}\n"
        f"Generated: {NOW[:10]} (seed data)\n"
    )
    (cdir / "F_Final" / "final_report.en.md").write_text(placeholder, encoding="utf-8")

    try:
        from docx import Document
        doc = Document()
        doc.add_heading(f"{workflow_type.replace('_', ' ').title()} — Test Report", 0)
        doc.add_paragraph(f"Client: {CLIENT_NAME}")
        doc.add_paragraph(f"Case ID: {case_id}")
        doc.add_paragraph(f"Generated: {NOW[:10]} (seed data)")
        doc.save(str(cdir / "F_Final" / "final_report.en.docx"))
    except ImportError:
        pass  # python-docx not installed — .md is sufficient for testing


def main() -> None:
    pm = ProjectManager()

    if pm.detect_slug_collision(SLUG):
        print(f"Engagement '{SLUG}' already exists — skipping creation.")
    else:
        intake = ProjectIntake(
            project_name=PROJECT_NAME,
            client_name=CLIENT_NAME,
            service_type="Investigation Report",
            language_standard="acfe",
        )
        pm.create_project(intake)
        print(f"Created engagement: cases/{SLUG}/")

    for workflow_type, case_id in CASE_IDS.items():
        _seed_case(workflow_type, case_id)
        print(f"Seeded case: {case_id} ({workflow_type})")

    # Register cases in the engagement's state.json
    from tools.project_manager import ProjectManager as _PM
    _pm = _PM()
    state = _pm.get_project(SLUG)
    if state is not None:
        state.cases = dict(CASE_IDS)
        _pm._write_state(SLUG, state)
        print(f"Updated engagement state with {len(CASE_IDS)} cases.")

    # Rebuild index so Case Tracker and Workspace see the new cases
    from tools.file_tools import build_case_index
    build_case_index()
    print("Index rebuilt.")
    print("\nSeed complete. Run: streamlit run app.py")


if __name__ == "__main__":
    main()
